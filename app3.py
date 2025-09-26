import streamlit as st
import openai
import os
from dotenv import load_dotenv
from docx import Document

# Load API key
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

st.title("📄 AI Resume Analyzer Pro")

# File uploader
resume_file = st.file_uploader("Upload your Resume (DOCX only)", type=["docx"])
job_desc = st.text_area("Paste Job Description here")

# Extra features
report_type = st.selectbox("Select Report Type", ["Short", "Medium", "Detailed"])

if st.button("Analyze Resume"):
    if resume_file and job_desc:
        # Read resume
        doc = Document(resume_file)
        resume_text = "\n".join([para.text for para in doc.paragraphs])

        # Prompt for ATS + Match
        prompt = f"""
        You are an ATS Resume Analyzer.
        Compare the following resume with the job description.
        Provide a {report_type} analysis including:
        - Match Score
        - Strengths
        - Weaknesses
        - Missing Keywords
        - ATS Optimization Suggestions
        Resume: {resume_text}
        Job Description: {job_desc}
        """

        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5
        )

        analysis = response.choices[0].message["content"]
        st.subheader("🔍 Resume Analysis")
        st.write(analysis)

        # Download Option
        def save_to_docx(text, filename="resume_report.docx"):
            doc = Document()
            doc.add_paragraph(text)
            doc.save(filename)
            return filename

        file_name = save_to_docx(analysis)
        with open(file_name, "rb") as f:
            st.download_button("📥 Download Report", f, file_name)
    else:
        st.warning("Please upload a resume and paste a job description.")
